from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
UML = DOCS / "uml"
OUT = DOCS / "Trabalho_TechFlow_Task_Manager.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)


def main():
    UML.mkdir(parents=True, exist_ok=True)
    use_case_png = UML / "casos-de-uso.png"
    class_png = UML / "diagrama-classes.png"
    make_use_case_diagram(use_case_png)
    make_class_diagram(class_png)

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_project_description(doc)
    add_methodology(doc)
    add_modeling_section(doc, use_case_png, class_png)
    add_quality_section(doc)
    add_change_section(doc)
    add_github_prints_section(doc)
    add_market_example(doc)
    add_final_notes(doc)
    add_references(doc)
    doc.save(OUT)
    print(OUT)


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "TechFlow Task Manager - Engenharia de Software"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Construindo um Projeto Ágil no GitHub")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Da Gestão ao Controle de Qualidade")
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_BLUE

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(24)
    run = meta.add_run(
        "Projeto acadêmico - TechFlow Solutions\n"
        "Sistema web de gerenciamento de tarefas para uma startup de logística"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Resumo executivo",
        "Este trabalho apresenta um projeto prático de Engenharia de Software com repositório Git, "
        "Kanban, CRUD web, testes automatizados, GitHub Actions e documentação de mudança de escopo."
    )

    doc.add_page_break()


def add_project_description(doc):
    doc.add_heading("1. Descrição do projeto e escopo inicial", level=1)
    add_body(
        doc,
        "A TechFlow Solutions foi contratada por uma startup de logística para desenvolver um sistema "
        "web de gerenciamento de tarefas. A solução permite registrar atividades, acompanhar responsáveis, "
        "visualizar o andamento do trabalho em um quadro Kanban e manter evidências do processo de "
        "desenvolvimento no GitHub."
    )
    add_body(
        doc,
        "O escopo inicial foi definido para manter o produto simples, funcional e alinhado aos objetivos "
        "da disciplina. O foco ficou no ciclo completo de software: planejamento, modelagem, versionamento, "
        "implementação, testes, automação e documentação."
    )
    add_bullets(
        doc,
        [
            "Cadastrar tarefas com título, descrição, responsável e prazo.",
            "Listar tarefas nas colunas A fazer, Em progresso e Concluído.",
            "Atualizar o status de cada tarefa conforme o avanço do trabalho.",
            "Excluir tarefas cadastradas incorretamente ou que não fazem mais parte do escopo.",
            "Filtrar tarefas por status e pesquisar por texto.",
            "Executar testes automatizados e validação de qualidade no GitHub Actions.",
        ],
    )

    doc.add_heading("Estrutura prática do repositório", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, [2.1, 4.2])
    set_header_cells(table.rows[0].cells, ["Diretório", "Finalidade"])
    rows = [
        ("src/", "Código da API, persistência e interface web."),
        ("tests/", "Testes automatizados do CRUD e validações."),
        ("docs/", "Documentação acadêmica, Kanban e diagramas UML."),
        (".github/workflows/", "Pipeline de integração contínua com GitHub Actions."),
        ("scripts/", "Scripts auxiliares de qualidade e geração do relatório."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right


def add_methodology(doc):
    doc.add_heading("2. Metodologia ágil utilizada", level=1)
    add_body(
        doc,
        "A metodologia adotada foi híbrida, combinando Kanban com práticas leves de Scrum. O Kanban "
        "organiza o fluxo nas colunas To Do, In Progress e Done, enquanto o Scrum contribui com backlog "
        "priorizado, incrementos curtos e revisão constante do que foi entregue."
    )
    add_bullets(
        doc,
        [
            "Kanban: visualização do trabalho, redução de tarefas esquecidas e acompanhamento do fluxo.",
            "Backlog incremental: divisão do projeto em entregas pequenas e verificáveis.",
            "Commits semânticos: histórico claro e rastreável das alterações.",
            "Integração contínua: execução automática de testes e qualidade a cada alteração publicada.",
        ],
    )

    doc.add_heading("Cards principais do Kanban", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, [0.6, 4.0, 1.4])
    set_header_cells(table.rows[0].cells, ["Nº", "Card", "Coluna"])
    cards = [
        ("1", "Definir escopo inicial", "Done"),
        ("2", "Criar repositório público", "Done"),
        ("3", "Implementar modelo de tarefas", "Done"),
        ("4", "Implementar API CRUD", "Done"),
        ("5", "Criar interface Kanban", "Done"),
        ("6", "Adicionar filtros básicos", "Done"),
        ("7", "Criar testes automatizados", "Done"),
        ("8", "Configurar GitHub Actions", "Done"),
        ("9", "Revisar documentação", "In Progress"),
        ("10", "Preparar vídeo pitch", "To Do"),
        ("11", "Adicionar prioridade para tarefas críticas", "Done"),
    ]
    for number, card, column in cards:
        cells = table.add_row().cells
        cells[0].text = number
        cells[1].text = card
        cells[2].text = column


def add_modeling_section(doc, use_case_png, class_png):
    doc.add_heading("3. Importância da modelagem", level=1)
    add_body(
        doc,
        "A modelagem é uma etapa essencial da Engenharia de Software porque transforma necessidades de "
        "negócio em representações compreensíveis antes da implementação. Ela reduz ambiguidades, melhora "
        "a comunicação entre estudantes, professores e stakeholders, e ajuda a identificar responsabilidades "
        "do sistema."
    )
    add_body(
        doc,
        "Neste projeto, os diagramas UML foram usados para representar as interações do usuário com o "
        "sistema e a estrutura principal das classes que sustentam o CRUD de tarefas."
    )

    doc.add_heading("3.1 Diagrama de Casos de Uso", level=2)
    add_picture(doc, use_case_png, "Figura 1 - Diagrama de Casos de Uso.")
    add_body(
        doc,
        "O diagrama mostra que gestores e membros da equipe interagem com funções como cadastrar, listar, "
        "filtrar, atualizar e excluir tarefas. O GitHub Actions aparece como ator externo responsável por "
        "executar testes e validar qualidade."
    )

    doc.add_heading("3.2 Diagrama de Classes", level=2)
    add_picture(doc, class_png, "Figura 2 - Diagrama de Classes.")
    add_body(
        doc,
        "O diagrama destaca a classe Task, que representa os dados de cada tarefa, e a classe TaskStore, "
        "responsável por persistência, criação, atualização, listagem e remoção. A camada ServerApi expõe "
        "essas operações para a interface web."
    )


def add_quality_section(doc):
    doc.add_heading("4. Testes automatizados e controle de qualidade", level=1)
    add_body(
        doc,
        "O controle de qualidade foi configurado com testes automatizados usando o módulo nativo node:test. "
        "Os testes cobrem o fluxo principal de criação, listagem, atualização, exclusão e validação de "
        "entradas inválidas."
    )
    add_body(
        doc,
        "Além dos testes, o projeto possui um script de qualidade que executa verificação de sintaxe nos "
        "arquivos JavaScript e identifica espaços em branco no final das linhas. O GitHub Actions executa "
        "essas validações automaticamente em pushes e pull requests."
    )
    add_callout(
        doc,
        "Resultado esperado",
        "A cada alteração enviada ao GitHub, o workflow CI deve baixar o código, configurar Node.js, "
        "instalar dependências, executar o lint e rodar os testes automatizados."
    )


def add_change_section(doc):
    doc.add_heading("5. Gestão de mudanças", level=1)
    add_body(
        doc,
        "Durante a simulação, o cliente solicitou uma mudança de escopo: destacar tarefas críticas para "
        "facilitar decisões em uma operação logística. A solicitação é coerente com o contexto do cliente, "
        "pois atrasos em tarefas críticas podem impactar entregas e atendimento ao consumidor final."
    )
    add_bullets(
        doc,
        [
            "Novo campo de prioridade no cadastro de tarefas.",
            "Valores disponíveis: Baixa, Média, Alta e Crítica.",
            "Filtro por prioridade na API e na interface web.",
            "Destaque visual para tarefas críticas.",
            "Atualização dos testes automatizados para cobrir o novo comportamento.",
        ],
    )
    add_body(
        doc,
        "A alteração foi registrada no Kanban como novo card e implementada em commit próprio, mantendo "
        "rastreabilidade entre justificativa de negócio, código, testes e documentação."
    )


def add_github_prints_section(doc):
    doc.add_heading("6. Prints comentados do GitHub", level=1)
    add_body(
        doc,
        "Os prints devem ser inseridos após publicar o repositório no GitHub e executar o workflow de CI. "
        "Abaixo estão os espaços e comentários já preparados para a versão final da entrega."
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, [2.1, 4.2])
    set_header_cells(table.rows[0].cells, ["Print necessário", "Comentário para a entrega"])
    rows = [
        (
            "Kanban no GitHub Projects",
            "O quadro evidencia a organização das tarefas nas colunas To Do, In Progress e Done, incluindo o card da mudança de escopo.",
        ),
        (
            "Histórico de commits",
            "Os commits semânticos demonstram evolução incremental e facilitam a rastreabilidade das decisões do projeto.",
        ),
        (
            "GitHub Actions aprovado",
            "O workflow comprova que testes e validação de qualidade foram automatizados no ciclo de desenvolvimento.",
        ),
    ]
    for label, comment in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = comment


def add_market_example(doc):
    doc.add_heading("7. Exemplo de solução existente", level=1)
    add_body(
        doc,
        "Um exemplo concreto de solução existente é o Trello, que usa quadros Kanban com cartões e colunas "
        "para organizar atividades. Outra solução comum em equipes de software é o Jira. O GitHub Projects "
        "cumpre papel semelhante, mas possui a vantagem de estar integrado ao repositório, às issues, aos "
        "pull requests, aos commits e ao GitHub Actions."
    )


def add_final_notes(doc):
    doc.add_heading("8. Considerações finais", level=1)
    add_body(
        doc,
        "O projeto demonstra como a Engenharia de Software integra planejamento, modelagem, versionamento, "
        "implementação, testes e adaptação a mudanças. A atividade também mostra que ferramentas como "
        "GitHub, GitHub Projects e GitHub Actions ajudam a reduzir falhas comuns em projetos ágeis, como "
        "baixa visibilidade do progresso, comunicação fragmentada e ausência de validação contínua."
    )


def add_references(doc):
    doc.add_heading("Referências", level=1)
    add_bullets(
        doc,
        [
            "GitHub Docs. GitHub Actions Documentation.",
            "Pressman, Roger S. Engenharia de Software: Uma Abordagem Profissional.",
            "Atlassian. Como usar Kanban para melhorar produtividade.",
            "Canal Programação Fácil. Testes automatizados com GitHub Actions.",
        ],
    )


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    paragraph.add_run(f"\n{body}")
    for paragraph in cell.paragraphs:
      paragraph.paragraph_format.space_after = Pt(4)


def add_picture(doc, image_path, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.25))

    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(10)
    caption_run = caption_paragraph.runs[0]
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = MUTED


def set_header_cells(cells, values):
    for cell, value in zip(cells, values):
        cell.text = value
        shade_cell(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def set_table_widths(table, widths):
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def make_use_case_diagram(path):
    image = Image.new("RGB", (1400, 860), "white")
    draw = ImageDraw.Draw(image)
    fonts = load_fonts()
    draw.rounded_rectangle((330, 90, 1070, 760), radius=24, outline="#2E74B5", width=4)
    draw.text((505, 110), "TechFlow Task Manager", fill="#1F4D78", font=fonts["title"])

    actors = {
        "Gestor de Projeto": (115, 245),
        "Membro da Equipe": (115, 520),
        "GitHub Actions": (1200, 400),
    }
    for label, (x, y) in actors.items():
        draw_actor(draw, x, y, label, fonts)

    cases = [
        ("Cadastrar tarefa", 500, 190),
        ("Listar tarefas", 800, 190),
        ("Atualizar status", 500, 330),
        ("Excluir tarefa", 800, 330),
        ("Filtrar tarefas", 500, 470),
        ("Priorizar tarefa crítica", 800, 470),
        ("Executar testes", 500, 620),
        ("Validar qualidade", 800, 620),
    ]
    case_positions = {}
    for label, x, y in cases:
        box = (x - 140, y - 42, x + 140, y + 42)
        draw.ellipse(box, fill="#F6F9FC", outline="#2E74B5", width=3)
        center_text(draw, label, box, fonts["body"], "#1B2430")
        case_positions[label] = (x, y)

    left_links = [
        ("Gestor de Projeto", "Cadastrar tarefa"),
        ("Gestor de Projeto", "Filtrar tarefas"),
        ("Gestor de Projeto", "Priorizar tarefa crítica"),
        ("Membro da Equipe", "Listar tarefas"),
        ("Membro da Equipe", "Atualizar status"),
        ("Membro da Equipe", "Excluir tarefa"),
    ]
    for actor, case in left_links:
        draw.line((actors[actor][0] + 85, actors[actor][1], case_positions[case][0] - 140, case_positions[case][1]), fill="#657184", width=2)

    for case in ["Executar testes", "Validar qualidade"]:
        draw.line((actors["GitHub Actions"][0] - 85, actors["GitHub Actions"][1], case_positions[case][0] + 140, case_positions[case][1]), fill="#657184", width=2)

    image.save(path)


def make_class_diagram(path):
    image = Image.new("RGB", (1400, 900), "white")
    draw = ImageDraw.Draw(image)
    fonts = load_fonts()

    task = draw_class(
        draw,
        (80, 120, 440, 515),
        "Task",
        [
            "id: String",
            "title: String",
            "description: String",
            "assignee: String",
            "dueDate: String",
            "priority: String",
            "status: String",
            "createdAt: String",
            "updatedAt: String",
        ],
        [],
        fonts,
    )
    store = draw_class(
        draw,
        (540, 110, 900, 520),
        "TaskStore",
        ["filePath: String"],
        ["list()", "create(input)", "update(id, changes)", "remove(id)", "clear()"],
        fonts,
    )
    api = draw_class(
        draw,
        (1010, 130, 1320, 500),
        "ServerApi",
        [],
        ["GET /api/tasks", "POST /api/tasks", "PUT /api/tasks/{id}", "DELETE /api/tasks/{id}"],
        fonts,
    )
    validation = draw_class(
        draw,
        (370, 640, 670, 800),
        "ValidationError",
        ["errors: Array"],
        [],
        fonts,
    )
    not_found = draw_class(
        draw,
        (760, 640, 1060, 800),
        "NotFoundError",
        ["id: String"],
        [],
        fonts,
    )

    draw_connector(draw, (task[2], 318), (store[0], 318), "0..*")
    draw_connector(draw, (store[2], 318), (api[0], 318), "usa")
    draw_connector(draw, (720, store[3]), (520, validation[1]), "valida")
    draw_connector(draw, (740, store[3]), (910, not_found[1]), "notifica")

    image.save(path)


def draw_actor(draw, x, y, label, fonts):
    draw.ellipse((x - 22, y - 70, x + 22, y - 26), outline="#1B2430", width=3)
    draw.line((x, y - 26, x, y + 42), fill="#1B2430", width=3)
    draw.line((x - 50, y, x + 50, y), fill="#1B2430", width=3)
    draw.line((x, y + 42, x - 42, y + 98), fill="#1B2430", width=3)
    draw.line((x, y + 42, x + 42, y + 98), fill="#1B2430", width=3)
    center_text(draw, label, (x - 120, y + 110, x + 120, y + 160), fonts["body"], "#1B2430")


def draw_class(draw, box, title, attributes, methods, fonts):
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill="#F6F9FC", outline="#2E74B5", width=3)
    draw.rectangle((x1, y1, x2, y1 + 58), fill="#E8EEF5", outline="#2E74B5", width=3)
    center_text(draw, title, (x1, y1, x2, y1 + 58), fonts["subtitle"], "#1F4D78")
    attr_y = y1 + 76
    for attr in attributes:
        draw.text((x1 + 22, attr_y), f"+ {attr}", fill="#1B2430", font=fonts["small"])
        attr_y += 32
    sep_y = max(attr_y + 8, y1 + 130)
    draw.line((x1, sep_y, x2, sep_y), fill="#2E74B5", width=2)
    method_y = sep_y + 18
    for method in methods:
        draw.text((x1 + 22, method_y), f"+ {method}", fill="#1B2430", font=fonts["small"])
        method_y += 32
    return box


def draw_connector(draw, start, end, label):
    draw.line((*start, *end), fill="#657184", width=3)
    mid = ((start[0] + end[0]) // 2, (start[1] + end[1]) // 2)
    draw.text((mid[0] - 24, mid[1] - 28), label, fill="#657184", font=load_fonts()["small"])


def center_text(draw, text, box, font, fill):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 20
    lines = []
    words = text.split()
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    line_height = font.size + 6
    total_height = line_height * len(lines)
    y = y1 + ((y2 - y1) - total_height) / 2
    for line in lines:
        width = draw.textlength(line, font=font)
        draw.text((x1 + ((x2 - x1) - width) / 2, y), line, fill=fill, font=font)
        y += line_height


def load_fonts():
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    font_path = next((candidate for candidate in candidates if candidate.exists()), None)

    if font_path is None:
        return {
            "title": ImageFont.load_default(),
            "subtitle": ImageFont.load_default(),
            "body": ImageFont.load_default(),
            "small": ImageFont.load_default(),
        }

    return {
        "title": ImageFont.truetype(str(font_path), 34),
        "subtitle": ImageFont.truetype(str(font_path), 28),
        "body": ImageFont.truetype(str(font_path), 24),
        "small": ImageFont.truetype(str(font_path), 22),
    }


if __name__ == "__main__":
    main()
